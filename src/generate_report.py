"""
generate_report.py
Auto-generates a Word report from pain point prioritization results.
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

def generate_word_report(priority_csv, pain_posts_csv, output_docx, top_n=5, posts_per_pain=3):
    # Create output directory if it doesn't exist
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    
    # Load data
    priority_df = pd.read_csv(priority_csv)
    pain_df = pd.read_csv(pain_posts_csv)
    
    # Calculate average metrics per keyword
    pain_metrics = pain_df.groupby('matched_keywords').agg({
        'upvotes': 'mean',
        'num_comments': 'mean'
    }).reset_index()
    pain_metrics['matched_keywords'] = pain_metrics['matched_keywords'].apply(eval)  # Convert string to list
    
    # Create document
    doc = Document()
    doc.add_heading('Pain Point Analysis Report', 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Total Posts Analyzed: {len(pain_df)}")
    doc.add_paragraph(f"Top Pain Points Identified: {top_n}")

    doc.add_heading('Summary Table', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pain Point'
    hdr_cells[1].text = 'Frequency'
    hdr_cells[2].text = 'Avg Upvotes'
    hdr_cells[3].text = 'Avg Comments'
    hdr_cells[4].text = 'Top Subreddits'

    # Add summary rows
    for _, row in priority_df.head(top_n).iterrows():
        keyword = row['keyword']
        freq = int(row['count'])
        
        # Find metrics for this keyword
        keyword_metrics = pain_metrics[pain_metrics['matched_keywords'].apply(lambda x: keyword in x)].agg({
            'upvotes': 'mean',
            'num_comments': 'mean'
        })
        
        avg_up = f"{keyword_metrics['upvotes']:.1f}" if not pd.isna(keyword_metrics['upvotes']) else "0.0"
        avg_com = f"{keyword_metrics['num_comments']:.1f}" if not pd.isna(keyword_metrics['num_comments']) else "0.0"
        
        # Find top subreddits for this keyword
        matching_posts = pain_df[pain_df['matched_keywords'].apply(lambda x: keyword in eval(x))]
        subreddits = matching_posts['subreddit'].value_counts().head(3)
        sub_list = ', '.join(subreddits.index)
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(keyword)
        row_cells[1].text = str(freq)
        row_cells[2].text = avg_up
        row_cells[3].text = avg_com
        row_cells[4].text = sub_list

    # Add detailed sections for each pain point
    for _, row in priority_df.head(top_n).iterrows():
        keyword = row['keyword']
        doc.add_heading(f"Pain Point: {keyword}", level=2)
        doc.add_paragraph(f"Frequency: {int(row['count'])}")
        
        # Find metrics for this keyword
        keyword_metrics = pain_metrics[pain_metrics['matched_keywords'].apply(lambda x: keyword in x)].agg({
            'upvotes': 'mean',
            'num_comments': 'mean'
        })
        
        doc.add_paragraph(f"Average Upvotes: {keyword_metrics['upvotes']:.1f}")
        doc.add_paragraph(f"Average Comments: {keyword_metrics['num_comments']:.1f}")
        
        # Find posts for this keyword
        matching_posts = pain_df[pain_df['matched_keywords'].apply(lambda x: keyword in eval(x))]
        subreddits = matching_posts['subreddit'].value_counts().head(3)
        doc.add_paragraph(f"Top Subreddits: {', '.join(subreddits.index)}")

        # Example posts
        doc.add_paragraph("Example Posts:", style='List Bullet')
        top_posts = matching_posts.sort_values('upvotes', ascending=False).head(posts_per_pain)
        for _, post in top_posts.iterrows():
            excerpt = post['body'][:200].replace('\n', ' ') + ('...' if len(post['body']) > 200 else '')
            doc.add_paragraph(f"• {post['title']} (r/{post['subreddit']}, {post['upvotes']} upvotes)\n  {excerpt}\n  [Link]({post['url']})", style='List Bullet')

    # Save document
    doc.save(output_docx)
    print(f"\n✓ Word report generated: {output_docx}")

# Example usage (uncomment and update filenames to run directly)
if __name__ == "__main__":
    # Example usage with data directory
    generate_word_report(
        priority_csv='data/pain_point_summary_20250430_032508.csv',
        pain_posts_csv='data/pain_point_posts_20250430_032508.csv',
        output_docx='data/Pain_Point_Analysis_Report.docx',
        top_n=5,
        posts_per_pain=3
    )